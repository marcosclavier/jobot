"""
Job Application Automation Bot

This script automates the process of finding and applying for jobs.
It now includes functionality to parse resumes, enhance user profiles using AI,
and manage the bot through a CLI interface.
"""

import os
import json
import time
import schedule
import requests
import logging
import click
import PyPDF2
import docx
import google.generativeai as genai
import hashlib
import base64
from cryptography.fernet import Fernet
from dotenv import load_dotenv
from datetime import datetime
from bs4 import BeautifulSoup
from concurrent.futures import ThreadPoolExecutor, as_completed, TimeoutError as FutureTimeoutError
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from google.api_core import exceptions as google_exceptions
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from docx.shared import Pt, Inches
import re
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Import submission logic from separate module
import submission

# --- Logging Configuration ---
logging.basicConfig(level=logging.INFO,
                    format='%(asctime)s - %(levelname)s - %(message)s',
                    handlers=[
                        logging.FileHandler("job_bot.log"),
                        logging.StreamHandler()
                    ])

# Load environment variables from .env file
load_dotenv()

# --- Constants ---
PROFILE_FILE = 'profile.json'
SEEN_JOBS_FILE = 'seen_jobs.json'
RECOMMENDED_JOBS_FILE = 'recommended_jobs.json'
SELECTED_JOBS_FILE = 'selected_jobs.json'
GENERATED_MATERIALS_FILE = 'generated_materials.json'
EDITED_MATERIALS_FILE = 'edited_materials.json'
PROFILE_HASH_FILE = '.profile_hash'
MAX_WORKERS = 5
SCRAPE_TIMEOUT = 30
MAX_PAGES = 3
MAX_DAYS_OLD = 15

# --- Encryption ---
def load_key():
    """Loads the encryption key from .env file."""
    key = os.getenv("ENCRYPTION_KEY")
    if not key:
        logging.error("ENCRYPTION_KEY not found in .env file. Generate one with `generate-key`.")
        raise ValueError("Encryption key not found.")
    return key.encode()

def encrypt_data(data: bytes, key):
    """Encrypts data using the provided key."""
    f = Fernet(key)
    return f.encrypt(data)

def decrypt_data(encrypted_data, key):
    """Decrypts data using the provided key."""
    f = Fernet(key)
    return f.decrypt(encrypted_data).decode()

from file_utils import load_json_file, save_json_file, load_profile

def save_profile(profile_data):
    """Encrypts and saves the user profile. Returns True on success, False on failure."""
    try:
        key = load_key()
        profile_json = json.dumps(profile_data, indent=4)
        encrypted_data = encrypt_data(profile_json, key)
        with open(PROFILE_FILE, 'wb') as f:
            f.write(encrypted_data)
        logging.info(f"Profile successfully encrypted and saved to {PROFILE_FILE}.")
        update_profile_hash()
        return True
    except (IOError, ValueError) as e:
        logging.error(f"Error saving profile: {e}", exc_info=True)
        return False

def validate_profile(profile):
    """Validates that the profile has essential fields."""
    required_fields = ['skills', 'location']
    missing_fields = [field for field in required_fields if not profile.get(field)]
    if missing_fields:
        logging.warning(f"Profile is incomplete. Missing fields: {', '.join(missing_fields)}.")
        return False
    return True

# --- Change Detection ---
def get_file_hash(file_path):
    """Computes the SHA256 hash of a file."""
    if not os.path.exists(file_path):
        return None
    hasher = hashlib.sha256()
    with open(file_path, 'rb') as f:
        buf = f.read()
        hasher.update(buf)
    return hasher.hexdigest()

def has_profile_changed():
    """Checks if profile.json has changed since the last run."""
    current_hash = get_file_hash(PROFILE_FILE)
    if not os.path.exists(PROFILE_HASH_FILE):
        return True  # Hash file doesn't exist, assume change

    with open(PROFILE_HASH_FILE, 'r') as f:
        stored_hash = f.read().strip()
    
    return current_hash != stored_hash

def update_profile_hash():
    """Updates the stored hash of the profile."""
    current_hash = get_file_hash(PROFILE_FILE)
    if current_hash:
        with open(PROFILE_HASH_FILE, 'w') as f:
            f.write(current_hash)
        logging.info("Updated profile hash.")

def sanitize_profile_for_generation(profile):
    """Remove personal info from profile copy to prevent inclusion in generated content."""
    sanitized = profile.copy()
    sanitized.pop('name', None)
    sanitized.pop('contact_info', None)
    sanitized.pop('location', None)  # Optional: Keep if needed for tailoring, but remove if not.
    return sanitized

# --- Resume Parsing ---
def extract_text_from_pdf(file_path):
    """Extracts text from a PDF file."""
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
        return text
    except Exception as e:
        logging.error(f"Failed to extract text from PDF {file_path}: {e}", exc_info=True)
        return None

def extract_text_from_docx(file_path):
    """Extracts text from a DOCX file."""
    try:
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        logging.error(f"Failed to extract text from DOCX {file_path}: {e}", exc_info=True)
        return None

def parse_resume(file_path):
    """Parses a resume file (PDF or DOCX) to extract text."""
    if not os.path.exists(file_path):
        logging.error(f"Resume file not found at: {file_path}")
        return None
    
    _, extension = os.path.splitext(file_path.lower())
    if extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif extension == '.docx':
        return extract_text_from_docx(file_path)
    else:
        logging.error(f"Unsupported file format: {extension}. Please use PDF or DOCX.")
        return None

# --- AI-Powered Content Generation ---
def enhance_profile_with_gemini(resume_text):
    """Uses Gemini to enhance the user profile based on resume text."""
    if not resume_text:
        logging.error("Cannot enhance profile, resume text is empty.")
        return None
    
    logging.info("Calling Gemini to enhance profile...")
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = f"""
        Analyze the following resume text and extract key information to create a professional profile.
        Based on the skills and experience, suggest a realistic salary range.
        Also, suggest a list of additional job search keywords (synonyms or related technologies).

        Extract the applicant's name, contact information (phone, email, LinkedIn if available).

        Resume Text:
        ---
        {resume_text}
        ---

        Return the output as a JSON object with the following keys:
        - "name": The applicant's full name.
        - "contact_info": A JSON object with "phone", "email", "linkedin" keys (use "" if not found).
        - "enhanced_skills": A list of key skills and technologies found.
        - "experience_summary": A brief summary of the professional experience.
        - "suggested_keywords": A list of 10-15 suggested keywords for job searching.
        - "salary_range": A suggested salary range (e.g., "$100,000 - $120,000 USD").

        Example JSON output:
        {{
          "name": "John Doe",
          "contact_info": {{"phone": "123-456-7890", "email": "john@example.com", "linkedin": "linkedin.com/in/johndoe"}},
          "enhanced_skills": ["Python", "Data Analysis", "Machine Learning", "SQL"],
          "experience_summary": "A data scientist with 5 years of experience...",
          "suggested_keywords": ["Data Scientist", "Python Developer", "AI Engineer"],
          "salary_range": "$110,000 - $135,000 USD"
        }}
        """
        response = model.generate_content(prompt)
        cleaned_response = response.text.strip().replace('```json', '').replace('```', '')
        logging.info("Successfully enhanced profile with Gemini.")
        return json.loads(cleaned_response)
    except google_exceptions.ResourceExhausted as e:
        logging.error(f"Gemini API quota exceeded: {e}", exc_info=True)
        return None
    except Exception as e:
        logging.error(f"Error enhancing profile with Gemini: {e}", exc_info=True)
        return None

def expand_keywords_with_gemini(base_keywords):
    """Expands a list of keywords using Gemini for better search results."""
    logging.info("Calling Gemini to expand keywords...")
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = f"""
        Given the following list of skills, generate a list of 20-30 related and synonymous keywords to improve job search results.
        Return only a comma-separated list of keywords.

        Skills: {', '.join(base_keywords)}
        """
        response = model.generate_content(prompt)
        expanded_keywords = [k.strip() for k in response.text.split(',')]
        logging.info("Successfully expanded keywords with Gemini.")
        return list(set(base_keywords + expanded_keywords)) # Combine and remove duplicates
    except google_exceptions.ResourceExhausted as e:
        logging.error(f"Gemini API quota exceeded: {e}", exc_info=True)
        return base_keywords
    except Exception as e:
        logging.error(f"Error expanding keywords with Gemini: {e}", exc_info=True)
        return base_keywords # Fallback to original keywords

def extract_questions_from_description(html_content):
    """Extracts potential application questions from a job description."""
    soup = BeautifulSoup(html_content, 'html.parser')
    questions = []
    for tag in soup.find_all(['li', 'p']):
        if tag.get_text(strip=True).endswith('?'):
            questions.append(tag.get_text(strip=True))
    return questions

def generate_application_materials(job_data, profile, custom_prompt=""):
    """Generates cover letter, resume suggestions, and answers to questions for a job."""
    job_title = job_data.get('job_details', {}).get('title', '')
    logging.info(f"Generating materials for: {job_title}")

    job_description = job_data.get('job_details', {}).get('full_description', '')
    questions = extract_questions_from_description(job_description)

    max_retries = 3
    for attempt in range(max_retries):
        try:
            model = genai.GenerativeModel('gemini-1.5-flash')
            
            prompt = (
                "**Objective:** Generate tailored application materials for a job applicant.\n\n"
                "**Applicant Profile:**\n"
                f"{json.dumps(sanitize_profile_for_generation(profile), indent=2)}\n\n"
                "**Job Details:**\n"
                f"{json.dumps(job_data['job_details'], indent=2)}\n\n"
                "**Custom Instructions:\n" +
                f"{custom_prompt}\n\n" + \
                "**Tasks:**\n"
                "1.  **Cover Letter:** Write a professional, enthusiastic, and tailored cover letter. It should highlight the applicant's most relevant skills and experiences from their profile that match the job description. Incorporate any custom instructions provided. Generate the body content only—do not include any headers, personal names, contact details (phone, email, LinkedIn), or location information, as these will be added separately. Use standard paragraph formatting. Do not use Markdown headings within the cover letter.\n"
                "2.  **Resume Adjustments:** Provide a list of specific, actionable suggestions for optimizing the applicant's resume for this job. Focus on incorporating keywords from the job description and aligning the experience summary with the role's requirements. Incorporate any custom instructions. Format each suggestion as a bullet point using `- ` or `* `. Use `**bold**` for keywords or important phrases.\n"
                "3.  **Answer Questions:** If there are questions below, provide thoughtful and detailed answers based on the applicant's profile. Format each answer as a standard paragraph.\n\n"
                "**Application Questions:**\n"
                f"{json.dumps(questions) if questions else 'No specific questions found.'}\n\n"
                "**Output Format:**\n"
                'Return a single JSON object with three keys: "cover_letter", "resume_suggestions", and "question_answers".\n'
                '- `cover_letter`: A string containing the full text of the cover letter.\n'
                '- `resume_suggestions`: A list of strings, where each string is a specific suggestion.\n'
                '- `question_answers`: A list of objects, each with "question" and "answer" keys.'
            )

            response = model.generate_content(prompt)
            # Attempt to extract JSON from potentially malformed response
            raw_response_text = response.text.strip()
            
            # Try to find the JSON block, assuming it might be wrapped in markdown
            if raw_response_text.startswith('```json'):
                cleaned_response = raw_response_text.replace('```json', '', 1).rsplit('```', 1)[0].strip()
            else:
                cleaned_response = raw_response_text.strip()

            generated_materials = json.loads(cleaned_response)
            job_data['generated_materials'] = generated_materials

            logging.info(f"Successfully generated materials for: {job_title}")
            return job_data

        except json.JSONDecodeError as e:
            logging.warning(f"JSONDecodeError on attempt {attempt + 1} for {job_title}: {e}. Raw response: {raw_response_text}")
            if attempt == max_retries - 1:
                logging.error(f"Failed to generate materials for {job_title} after {max_retries} attempts due to JSON decoding error.")
                return None
        except google_exceptions.ResourceExhausted as e:
            logging.error(f"Gemini API quota exceeded while generating materials for {job_title}: {e}", exc_info=True)
            return None
        except Exception as e:
            logging.error(f"Error generating materials for {job_title} on attempt {attempt + 1}: {e}", exc_info=True)
            if attempt == max_retries - 1:
                return None
    return None

# --- Material Refinement ---
def simulate_ats_score(job_data, materials):
    """Simulates an ATS score by matching keywords."""
    logging.info("Simulating ATS score...")
    job_details = job_data.get('job_details', {})
    job_description = job_details.get('full_description', '')
    cover_letter = materials.get('cover_letter', '')
    resume_text = materials.get('refined_resume', ' '.join(materials.get('resume_suggestions', [])))
    full_text = cover_letter + ' ' + resume_text

    if not job_description or not full_text:
        return {"score": 0, "matched_keywords": [], "missing_keywords": []}

    # Check if keywords are already cached
    if 'ats_keywords' in job_data:
        keywords = job_data['ats_keywords']
        logging.info("Using cached ATS keywords.")
    else:
        try:
            model = genai.GenerativeModel('gemini-1.5-flash')
            prompt = f"""
            From the following job description, extract the 20 most important keywords and skills.
            Return them as a JSON list of strings.

            Job Description:
            ---
            {job_description[:4000]}
            ---
            """
            response = model.generate_content(prompt)
            cleaned_response = response.text.strip().replace('```json', '').replace('```', '')
            keywords = json.loads(cleaned_response)
            job_data['ats_keywords'] = keywords  # Cache the extracted keywords
            logging.info("Extracted and cached new ATS keywords.")
        except Exception as e:
            logging.error(f"Error extracting ATS keywords: {e}", exc_info=True)
            return {"score": 0, "matched_keywords": [], "missing_keywords": []}

    matched_keywords = {kw for kw in keywords if kw.lower() in full_text.lower()}
    missing_keywords = set(keywords) - matched_keywords
    score = (len(matched_keywords) / len(keywords) * 100) if keywords else 0

    logging.info(f"ATS simulation complete. Score: {score:.2f}%")
    return {
        "score": round(score, 2),
        "matched_keywords": list(matched_keywords),
        "missing_keywords": list(missing_keywords)
    }


def validate_materials_with_gemini(materials):
    """Uses Gemini to validate the completeness and quality of application materials."""
    logging.info("Calling Gemini to validate materials...")
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = f"""
        Please act as a career coach and review the following application materials.
        Provide a list of actionable suggestions for improvement.
        Focus on clarity, impact, and professionalism. Check for any incomplete sentences or sections.
        If "refined_resume" is present, review that full resume text; otherwise, review "resume_suggestions".
        Do not comment on the absence of personal information (name, contact details, location) as these are intentionally excluded from the generated content.

        Materials:
        ---
        {json.dumps(materials, indent=2)}
        ---

        Return a JSON object with a single key "validation_feedback", which is a list of strings.
        Example: {{"validation_feedback": ["The cover letter could be more specific about project X.", "Consider rephrasing the second resume suggestion for more impact."]}}
        """
        response = model.generate_content(prompt)
        cleaned_response = response.text.strip().replace('```json', '').replace('```', '')
        validation = json.loads(cleaned_response)
        logging.info("Successfully validated materials with Gemini.")
        return validation.get('validation_feedback', [])
    except Exception as e:
        logging.error(f"Error validating materials with Gemini: {e}", exc_info=True)
        return ["Failed to get validation from Gemini."]

def get_gemini_suggestions(text_to_improve, instruction):
    """Uses Gemini to improve a specific piece of text based on instructions."""
    logging.info("Calling Gemini for text improvement suggestions...")
    try:
        model = genai.GenerativeModel('gemini-1.5-flash') # Using pro for better writing
        prompt = f"""
        Please revise the following text based on the user's instruction.
        Return only the revised text, without any extra formatting or explanation.

        Original Text:
        ---
        {text_to_improve}
        ---

        Instruction: "{instruction}"
        """
        response = model.generate_content(prompt)
        logging.info("Successfully received suggestion from Gemini.")
        return response.text.strip()
    except Exception as e:
        logging.error(f"Error getting suggestions from Gemini: {e}", exc_info=True)
        return "Failed to get suggestion from Gemini."

def apply_validation_feedback(materials, validation_feedback):
    """Uses Gemini to apply validation feedback to the materials."""
    if not validation_feedback:
        logging.info("No validation feedback to apply.")
        return materials

    logging.info("Applying validation feedback with Gemini...")
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = f"""
        Current Materials:
        ---
        {json.dumps(materials, indent=2)}
        ---

        Validation Feedback:
        ---
        {json.dumps(validation_feedback, indent=2)}
        ---

        Revise the "cover_letter" and "refined_resume" to address all the feedback points. Make sure the revisions improve clarity, impact, professionalism, and incorporate any suggestions. While revising, ensure the content remains optimized for ATS by incorporating key job description keywords without keyword stuffing.

        Return a JSON object with keys "cover_letter" and "refined_resume", containing the revised texts.
        """
        response = model.generate_content(prompt)
        cleaned_response = response.text.strip().replace('```json', '').replace('```', '')
        revised = json.loads(cleaned_response)
        materials['cover_letter'] = revised.get('cover_letter', materials['cover_letter'])
        materials['refined_resume'] = revised.get('refined_resume', materials['refined_resume'])
        logging.info("Successfully applied feedback.")
        return materials
    except Exception as e:
        logging.error(f"Error applying validation feedback: {e}", exc_info=True)
        return materials

# --- New Function for Refined Resume ---
def generate_refined_resume(profile, materials, job_data):
    """Uses Gemini to generate a refined resume based on profile and suggestions."""
    suggestions = materials.get('resume_suggestions', [])
    job_description = job_data.get('job_details', {}).get('full_description', '')
    job_title = job_data.get('job_details', {}).get('title', '')
    
    logging.info(f"Generating refined resume for: {job_title}")
    
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = f"""
        Applicant Profile:
        ---
        {json.dumps(profile, indent=2)}
        ---

        Job Description:
        ---
        {job_description[:4000]}
        ---

        Resume Suggestions:
        ---
        {json.dumps(suggestions, indent=2)}
        ---

        Task: Create a professional resume based on the applicant profile. Generate the body content only—do not include any headers, personal names, contact details (phone, email, LinkedIn), or location information, as these will be added separately. Incorporate all the provided suggestions. Tailor it specifically to the job description, highlighting relevant skills and experiences. Ensure the resume is professional, concise, and optimized for ATS systems by including keywords from the job description.
        
        Use the following formatting guidelines:
        - Use `##` for main section headings (e.g., `## Summary`, `## Experience`, `## Education`, `## Skills`).
        - For job/education entries, use the format `**Job Title, Company** (Dates)` or `**Degree, Institution** (Dates)`. Ensure the dates are enclosed in parentheses and italicized.
        - Use `- ` or `* ` for bullet points.
        - For skills lists, provide them as a comma-separated list if there are many, or as bullet points if there are fewer than 4.
        - Use `**bold**` and `*italic*` for emphasis where appropriate.
        
        Return only the full text of the resume, without any additional explanations or formatting.
        """
        response = model.generate_content(prompt)
        refined_resume = response.text.strip()
        logging.info(f"Successfully generated refined resume for: {job_title}")
        return refined_resume
    except google_exceptions.ResourceExhausted as e:
        logging.error(f"Gemini API quota exceeded while generating refined resume for {job_title}: {e}", exc_info=True)
        return "Failed to generate refined resume due to API quota."
    except Exception as e:
        logging.error(f"Error generating refined resume for {job_title}: {e}", exc_info=True)
        return "Failed to generate refined resume."



def save_recommended_job(job_data):
    recommendations = load_json_file(RECOMMENDED_JOBS_FILE)
    recommendations.append(job_data)
    save_json_file(recommendations, RECOMMENDED_JOBS_FILE)

def save_selected_job(job_data):
    selections = load_json_file(SELECTED_JOBS_FILE)
    selections.append(job_data)
    save_json_file(selections, SELECTED_JOBS_FILE)

def filter_new_jobs(jobs, seen_job_ids):
    """Filters out jobs that have already been seen."""
    return [job for job in jobs if job.get('id') and job.get('id') not in seen_job_ids]

import re  # Already imported, but ensure for regex

def scrape_full_description(job):
    redirect_url = job.get('redirect_url')
    original_description = job.get('description', '')

    if not redirect_url:
        job['full_description'] = original_description
        job['ats_url'] = None
        job['ats_platform'] = 'Unknown'
        return job

    session = requests.Session()
    retry = Retry(total=3, backoff_factor=0.5, status_forcelist=[500, 502, 503, 504])
    adapter = HTTPAdapter(max_retries=retry)
    session.mount('https://', adapter)
    session.mount('http://', adapter)

    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/108.0.0.0 Safari/537.36',
            'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8',
            'Accept-Language': 'en-US,en;q=0.9',
            'Accept-Encoding': 'gzip, deflate, br',
            'DNT': '1',
            'Upgrade-Insecure-Requests': '1'
        }
        page_response = session.get(redirect_url, headers=headers, timeout=15)
        page_response.raise_for_status()
        soup = BeautifulSoup(page_response.text, 'html.parser')
        
        # Existing description scraping
        description_div = soup.find('section', class_='adp-body') or soup.find('div', class_='job-description')
        job['full_description'] = description_div.get_text(separator='\n', strip=True) if description_div else original_description
        
        # New: Scrape Apply URL
        apply_a = soup.find('a', attrs={'data-js': 'apply'})  # Primary matcher from your snippet
        if not apply_a:
            # Fallback: Search by text or class
            apply_a = soup.find('a', string=re.compile('Apply for this job', re.I)) or \
                      soup.find('a', class_=re.compile('adzuna-green'))
        
        if apply_a:
            apply_href = apply_a['href']
            if apply_href.startswith('/'):  # Relative URL
                apply_href = 'https://www.adzuna.com' + apply_href
            
            # Follow redirects to get final ATS URL
            apply_response = session.get(apply_href, headers=headers, allow_redirects=True, timeout=15)
            final_url = apply_response.url
            job['ats_url'] = final_url
            
            # Detect ATS platform
            if 'greenhouse.io' in final_url:
                job['ats_platform'] = 'Greenhouse'
            elif 'lever.co' in final_url:
                job['ats_platform'] = 'Lever'
            elif 'bamboohr.com' in final_url:
                job['ats_platform'] = 'BambooHR'
            # Add more (e.g., 'workable.com': 'Workable', 'ashbyhq.com': 'Ashby')
            else:
                job['ats_platform'] = 'Unknown'
            logging.info(f"Scraped ATS URL: {final_url} ({job['ats_platform']})")
        else:
            job['ats_url'] = None
            job['ats_platform'] = 'Unknown'
            logging.warning("No Apply button found on page.")
    
    except requests.exceptions.RequestException as e:
        logging.error(f"Failed to fetch or process {redirect_url}: {e}")
        job['full_description'] = original_description
        job['ats_url'] = None
        job['ats_platform'] = 'Unknown'
        
    return job

def fetch_adzuna_jobs(profile, primary_keyword, secondary_keywords):
    """Fetches job listings from Adzuna with a primary and secondary keyword strategy."""
    app_id = os.getenv("ADZUNA_APP_ID")
    app_key = os.getenv("ADZUNA_APP_KEY")

    if not app_id or not app_key:
        logging.error("Adzuna API credentials not found.")
        return []

    base_url = "https://api.adzuna.com/v1/api/jobs/us/search"
    params = {
        'app_id': app_id,
        'app_key': app_key,
        'results_per_page': 50,
        'what': primary_keyword,
        'what_or': ' '.join(secondary_keywords),
        'where': profile.get('location', 'remote'),
        'sort_by': 'date',
        'max_days_old': MAX_DAYS_OLD
    }

    all_jobs = []
    for page in range(1, MAX_PAGES + 1):
        try:
            response = requests.get(f"{base_url}/{page}", params=params)
            response.raise_for_status()
            data = response.json()
            jobs = data.get('results', [])
            for job in jobs:
                job['source'] = 'adzuna'
            all_jobs.extend(jobs)
            if len(jobs) < params['results_per_page']:
                break
        except (requests.exceptions.RequestException, json.JSONDecodeError) as e:
            logging.error(f"Error fetching from Adzuna on page {page}: {e}", exc_info=True)
            break

    return all_jobs

def fetch_indeed_jobs(profile, keywords):
    """Fetches job listings from Indeed API with retries."""
    api_key = os.getenv("INDEED_API_KEY")
    if not api_key:
        logging.warning("Indeed API key not found, skipping Indeed search.")
        return []

    base_url = "https://api.indeed.com/ads/apisearch"
    params = {
        "publisher": api_key,
        "q": ' '.join(keywords),
        "l": profile.get('location', 'remote'),
        "sort": "date",
        "limit": 50,
        "v": "2",
        "format": "json",
        "jt": profile.get('work_type', 'fulltime'),
    }

    all_jobs = []
    session = requests.Session()
    retry = Retry(total=3, backoff_factor=1, status_forcelist=[429, 500, 502, 503, 504])
    adapter = HTTPAdapter(max_retries=retry)
    session.mount('https://', adapter)

    try:
        response = session.get(base_url, params=params, timeout=15)
        response.raise_for_status()
        data = response.json()
        for item in data.get('results', []):
            job = {
                'id': item.get('jobkey'),
                'title': item.get('jobtitle'),
                'description': item.get('snippet'),
                'redirect_url': item.get('url'),
                'company': {'display_name': item.get('company')},
                'source': 'indeed'
            }
            all_jobs.append(job)
    except (requests.exceptions.RequestException, json.JSONDecodeError) as e:
        logging.error(f"Error fetching from Indeed: {e}", exc_info=True)

    return all_jobs

def evaluate_job_fit(job, user_profile):
    """Evaluates job fit using Gemini and calculates skill match."""
    job_description = job.get('full_description', job.get('description', ''))
    user_skills = set(skill.lower() for skill in user_profile.get('skills', []))
    
    found_skills = {skill for skill in user_skills if skill in job_description.lower()}
    skill_match_percentage = (len(found_skills) / len(user_skills) * 100) if user_skills else 0

    logging.info(f"Calling Gemini to evaluate job fit for: {job.get('title', '')}")
    try:
        model = genai.GenerativeModel('gemini-1.5-flash')
        prompt = f"""
        User Profile: {json.dumps(sanitize_profile_for_generation(user_profile), indent=2)}
        Job Description: {job_description}

        Based on the user profile and job description, perform the following tasks:
        1. Rate the job fit on a scale of 1-10.
        2. Provide a brief explanation for the rating.
        3. Write a concise one-paragraph summary of the job role and its key responsibilities.

        Return a JSON object with the keys: "fit_score", "explanation", and "summary".
        Example: {"fit_score": 8, "explanation": "The role aligns well...", "summary": "This is a software engineering role..."}
        """
        response = model.generate_content(prompt)
        cleaned_response = response.text.strip().replace('```json', '').replace('```', '')
        evaluation = json.loads(cleaned_response)
        
        evaluation['skill_match_percentage'] = round(skill_match_percentage, 2)
        evaluation['matched_skills'] = list(found_skills)
        
        logging.info(f"Successfully evaluated job fit for: {job.get('title', '')}")
        return evaluation
    except google_exceptions.ResourceExhausted as e:
        logging.error(f"Gemini API quota exceeded while evaluating job fit: {e}", exc_info=True)
        return None
    except Exception as e:
        logging.error(f"Error evaluating job fit with Gemini: {e}", exc_info=True)
        return None

# --- CLI Command Group ---
@click.group()
def cli():
    """Job Application Bot CLI."""
    pass

@cli.command()
def search():
    """Fetches, filters, and evaluates new job listings from multiple sources."""
    logging.info("--- Running job search ---")
    profile = load_profile()
    if not validate_profile(profile):
        return

    if has_profile_changed():
        logging.info("Profile has changed. Clearing seen jobs to re-evaluate all.")
        save_json_file([], SEEN_JOBS_FILE) # Reset seen jobs
        update_profile_hash()

    all_skills = profile.get('skills', [])
    suggested_keywords = profile.get('suggested_keywords', [])
    primary_keyword = all_skills[0] if all_skills else ''
    secondary_keywords = all_skills[1:] + suggested_keywords

    safe_secondary_keywords = []
    current_length = 0
    for keyword in secondary_keywords:
        if current_length + len(keyword) + 1 > 512:
            break
        safe_secondary_keywords.append(keyword)
        current_length += len(keyword) + 1

    logging.info(f"Primary keyword: '{primary_keyword}'. Using {len(safe_secondary_keywords)} secondary keywords.")

    seen_job_ids = load_json_file(SEEN_JOBS_FILE)
    
    adzuna_jobs = fetch_adzuna_jobs(profile, primary_keyword, safe_secondary_keywords)
    indeed_keywords = [primary_keyword] + safe_secondary_keywords
    indeed_jobs = fetch_indeed_jobs(profile, indeed_keywords)
    all_jobs = adzuna_jobs + indeed_jobs

    if all_jobs:
        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
            all_jobs = list(executor.map(scrape_full_description, all_jobs))

        new_jobs = filter_new_jobs(all_jobs, seen_job_ids)
        if new_jobs:
            logging.info(f"Found {len(new_jobs)} new jobs to evaluate.")
            for job in new_jobs:
                title = job.get('title', '')
                company = job.get('company', {}).get('display_name', '')
                logging.info(f"--- Evaluating: {title} at {company} ---")
                
                evaluation = evaluate_job_fit(job, profile)
                if not evaluation:
                    logging.error(f"Could not evaluate job: {title}")
                    continue

                fit_score = evaluation.get('fit_score', 0)
                logging.info(f"  - Fit Score: {fit_score}/10")
                logging.info(f"  - Skill Match: {evaluation.get('skill_match_percentage')}% ")

                if fit_score >= 7:
                    logging.info(f"  - Recommendation: Saving job - {title}")
                    save_recommended_job({"job_details": job, "evaluation": evaluation})
                else:
                    logging.warning(f"  - Recommendation: Skipping job due to low fit score.")
            
            seen_job_ids.extend(job['id'] for job in new_jobs if job.get('id'))
        else:
            logging.info("No new jobs found.")
    else:
        logging.warning("No jobs found from any source.")

    save_json_file(seen_job_ids, SEEN_JOBS_FILE)
    logging.info("--- Job search finished ---")

@cli.command()
def review():
    """Interactively review recommended jobs."""
    recommendations = load_json_file(RECOMMENDED_JOBS_FILE)
    if not recommendations:
        logging.info("No recommended jobs to review. Run `search` first.")
        return

    keep_jobs, selected_jobs = [], []
    for job_data in recommendations:
        job_details = job_data.get('job_details', {})
        evaluation = job_data.get('evaluation', {})

        click.echo("\n" + "-"*80)
        click.echo(f"Title: {job_details.get('title', '')}")
        click.echo(f"Company: {job_details.get('company', {}).get('display_name', '')}")
        click.echo(f"Fit Score: {evaluation.get('fit_score')}/10 | Skill Match: {evaluation.get('skill_match_percentage')}% ")
        click.echo(f"Summary: {evaluation.get('summary')}")
        click.echo("-"*80)

        choice = click.prompt("Action (i=Interested, n=Not Interested, s=Save for later)", type=str, default='s').lower()

        if choice == 'i':
            selected_jobs.append(job_data)
            logging.info(f"Marked as interested: {job_details.get('title', '')}")
        elif choice == 'n':
            logging.info(f"Marked as not interested: {job_details.get('title', '')}.")
        else:
            keep_jobs.append(job_data)

    if selected_jobs:
        for job in selected_jobs:
            save_selected_job(job)
    save_json_file(keep_jobs, RECOMMENDED_JOBS_FILE)
    logging.info("Review session finished.")

@cli.command()
def generate():
    """Interactively generates application materials for selected jobs."""
    selected_jobs = load_json_file(SELECTED_JOBS_FILE)
    if not selected_jobs:
        logging.info("No selected jobs to generate materials for. Run `review` first.")
        return

    profile = load_profile()
    all_generated_materials = []
    
    for job_data in selected_jobs:
        job_title = job_data.get('job_details', {}).get('title', '')
        click.echo("\n" + "="*80)
        click.echo(f"Preparing to generate materials for: {job_title}")
        click.echo("="*80)

        # Initial generation
        generated_job_data = generate_application_materials(job_data, profile)
        if not generated_job_data:
            continue # Skip if initial generation fails

        # Automatically accept generated materials
        all_generated_materials.append(generated_job_data)
        logging.info(f"Automatically accepted materials for: {job_title}")
    
    save_json_file(all_generated_materials, GENERATED_MATERIALS_FILE)
    if all_generated_materials:
        logging.info(f"Finished generating materials for {len(all_generated_materials)} jobs.")
        click.echo(f"Run `refine` to edit and approve them, then `export-docs` to create DOCX files.")
    else:
        logging.info("No materials were generated in this session.")

@cli.command()
def refine():
    """Automatically refines, validates, and approves generated application materials."""
    generated_materials = load_json_file(GENERATED_MATERIALS_FILE)
    if not generated_materials:
        click.echo("No generated materials to refine. Run `generate` first.")
        logging.info("No generated materials to refine. Run `generate` first.")
        return

    approved_materials = []
    profile = load_profile()  # Load profile for refined resume generation
    
    for i, item in enumerate(generated_materials):
        job_details = item.get('job_details', {})
        materials = item.get('generated_materials', {}).copy() # Use a copy to edit
        job_data = item  # For refined resume
        job_title = job_details.get('title', '')

        click.echo(f"Refining {job_title}...")

        # Generate initial refined resume
        refined_resume = generate_refined_resume(profile, materials, item)
        materials['refined_resume'] = refined_resume

        # Simulate ATS
        ats_results = simulate_ats_score(item, materials)
        item['ats_score'] = ats_results

        # Validate
        validation_feedback = validate_materials_with_gemini(materials)

        # Apply feedback
        materials = apply_validation_feedback(materials, validation_feedback)

        # Re-simulate ATS after revision
        ats_results = simulate_ats_score(item, materials)
        item['ats_score'] = ats_results

        item['generated_materials'] = materials
        approved_materials.append(item)

    if approved_materials:
        save_json_file(approved_materials, EDITED_MATERIALS_FILE)
        click.echo(f"Saved {len(approved_materials)} approved material sets to {EDITED_MATERIALS_FILE}.")
    else:
        click.echo("No materials were approved in this session.")


def add_styled_text(paragraph, text):
    """Manually parses text for **bold** and *italic*, skipping unclosed markers."""
    i = 0
    while i < len(text):
        if text[i:i+2] == '**':
            # Look for closing **
            j = text.find('**', i + 2)
            if j != -1:
                # Add bold run for content between
                run = paragraph.add_run(text[i + 2:j])
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Force black
                i = j + 2
            else:
                # Unclosed **, skip the marker to remove it
                i += 2
        elif text[i] == '*':
            # Look for closing *
            j = text.find('*', i + 1)
            if j != -1:
                # Add italic run for content between
                run = paragraph.add_run(text[i + 1:j])
                run.italic = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Force black
                i = j + 1
            else:
                # Unclosed *, skip the marker
                i += 1
        else:
            # Add plain text until next marker
            next_bold = text.find('**', i)
            next_italic = text.find('*', i)
            next_mark = min(next_bold if next_bold != -1 else len(text),
                            next_italic if next_italic != -1 else len(text))
            run = paragraph.add_run(text[i:next_mark])
            run.font.color.rgb = RGBColor(0, 0, 0)  # Force black
            i = next_mark

def add_formatted_content(doc, content, is_resume=False):
    """Adds content to DOCX with native formatting, parsing Markdown and resume structure."""
    import re  # Ensure re is imported at the top of the script if not already

    lines = content.split('\n')
    i = 0
    is_skills = False
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        # Markdown headings
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)  # Force black
            is_skills = False
            i += 1
            continue
        elif line.startswith('## '):
            heading_text = line[3:]
            heading = doc.add_heading(heading_text, level=2)
            heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)  # Force black
            is_skills = (heading_text == 'Skills')
            i += 1
            continue
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
            heading.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)  # Force black
            is_skills = False
            i += 1
            continue
        # Job/education entries
        if is_resume and '(' in line and ')' in line and line.endswith(')'):
            date_start = line.rfind('(')
            title_part = line[:date_start].strip()
            date_part = line[date_start:].strip()
            p = doc.add_paragraph()
            add_styled_text(p, title_part)
            p.add_run(' ')
            run = p.add_run(date_part)
            run.italic = True
            run.font.color.rgb = RGBColor(0, 0, 0)  # Force black
            i += 1
            continue
        # Bullets
        elif line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            add_styled_text(p, line[2:].strip())
            i += 1
            continue
        # Skills subheadings and grouped items
        elif is_resume and is_skills:
            clean_line = re.sub(r'\*', '', line).strip()
            if clean_line.endswith(':'):
                # Add subheading
                p = doc.add_paragraph()
                add_styled_text(p, line)
                i += 1
                # Collect items for comma-separated list
                items = []
                while i < len(lines):
                    next_line = lines[i].strip()
                    if not next_line:
                        i += 1
                        continue
                    clean_next = re.sub(r'\*', '', next_line).strip()
                    if clean_next.endswith(':') or next_line.startswith('#') or next_line.startswith('- ') or next_line.startswith('* ') or next_line.endswith(')') or '.' in next_line:
                        break
                    if len(next_line) < 60:  # Likely a skill item
                        items.append(next_line)
                    i += 1
                if items:
                    joined = ', '.join([item.strip() for item in items if item.strip()])
                    p = doc.add_paragraph()
                    add_styled_text(p, joined)
                continue
        # Default paragraph
        p = doc.add_paragraph()
        add_styled_text(p, line)
        i += 1

def set_document_styles(doc):
    """Sets professional document-wide styles."""
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'  # Or 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)  # Force black
    
    # Force black for headings
    for i in range(1, 10):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.color.rgb = RGBColor(0, 0, 0)
    
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

def add_styled_header(doc, header_text, is_resume=False):
    """Adds a formatted header with name as heading, contact smaller."""
    lines = header_text.split('\n')
    
    if lines:  # Name as heading
        name_heading = doc.add_heading(lines[0], level=1)
        name_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if is_resume else WD_PARAGRAPH_ALIGNMENT.LEFT
        for run in name_heading.runs:
            run.font.color.rgb = RGBColor(0, 0, 0)  # Force black
    
    if len(lines) > 1:  # Contact line
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER if is_resume else WD_PARAGRAPH_ALIGNMENT.LEFT
        contact_parts = lines[1].split(' | ')
        for i, part in enumerate(contact_parts):
            run = contact_p.add_run(part.strip())
            if 'http' in part or 'linkedin' in part.lower():  # Italicize links
                run.italic = True
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)  # Force black
            if i < len(contact_parts) - 1:
                contact_p.add_run(' | ')
    
    # Optional: Add a thin separator line
    separator_p = doc.add_paragraph('___________________________________________')
    separator_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in separator_p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)  # Force black

@cli.command()
def export_docs():
    """Exports approved and edited materials to DOCX and JSON files."""
    approved_materials = load_json_file(EDITED_MATERIALS_FILE)
    if not approved_materials:
        click.echo("No approved materials to export. Run `refine` first.")
        logging.info("No approved materials to export. Run `refine` first.")
        return

    output_dir = "applications"
    os.makedirs(output_dir, exist_ok=True)

    profile = load_profile()
    name = profile.get('name', '')
    contact_info = profile.get('contact_info', {})
    location = profile.get('location', '')

    # Construct the personal information header
    header_text = f"{name}\n"
    contact_parts = []
    if contact_info.get('phone'):
        contact_parts.append(contact_info['phone'])
    if contact_info.get('email'):
        contact_parts.append(contact_info['email'])
    if contact_info.get('linkedin'):
        contact_parts.append(contact_info['linkedin'])
    header_text += " | ".join(contact_parts)
    if location:
        header_text += f"\n{location}"

    for item in approved_materials:
        job_details = item.get('job_details', {})
        materials = item.get('generated_materials', {})
        company = job_details.get('company', {}).get('display_name', 'N/A').replace(' ', '_').replace('/', '_')
        title = job_details.get('title', 'N/A').replace(' ', '_').replace('/', '_')

        # --- Export Cover Letter ---
        cover_letter_doc = docx.Document()
        set_document_styles(cover_letter_doc)
        
        # Add styled header (left-aligned for cover letters)
        add_styled_header(cover_letter_doc, header_text)
        
        # Add salutation
        salutation_p = cover_letter_doc.add_paragraph()
        add_styled_text(salutation_p, f"Dear hiring manager at {job_details.get('company', {}).get('display_name', 'N/A')},")
        
        # Add formatted cover letter body
        cover_letter_content = materials.get('cover_letter', '')
        add_formatted_content(cover_letter_doc, cover_letter_content)
        
        # Add closing
        closing_p = cover_letter_doc.add_paragraph()
        add_styled_text(closing_p, "Regards,")
        
        # Add two line breaks (empty paragraphs)
        cover_letter_doc.add_paragraph()
        cover_letter_doc.add_paragraph()
        
        # Add candidate name
        name_p = cover_letter_doc.add_paragraph()
        add_styled_text(name_p, name)
        
        cover_letter_path = os.path.join(output_dir, f"{company}_{title}_CoverLetter.docx")
        cover_letter_doc.save(cover_letter_path)
        logging.info(f"Exported cover letter to {cover_letter_path}")

        # --- Export Refined Resume ---
        resume_doc = docx.Document()
        set_document_styles(resume_doc)
        
        # Add styled header (center-aligned for resumes)
        add_styled_header(resume_doc, header_text, is_resume=True)
        
        # Add formatted resume body (with resume-specific parsing)
        resume_content = materials.get('refined_resume', '')
        add_formatted_content(resume_doc, resume_content, is_resume=True)
        
        resume_path = os.path.join(output_dir, f"{company}_{title}_Resume.docx")
        resume_doc.save(resume_path)
        logging.info(f"Exported refined resume to {resume_path}")

        # --- Export Question Answers (JSON) ---
        question_answers = materials.get('question_answers', [])
        if question_answers:
            qa_path = os.path.join(output_dir, f"{company}_{title}_Questions.json")
            save_json_file(question_answers, qa_path)
            logging.info(f"Exported questions and answers to {qa_path}")
        else:
            logging.info(f"No questions and answers to export for {company}_{title}.")

    click.echo(f"Successfully exported materials for {len(approved_materials)} jobs to the 'applications' folder.")

@cli.command()
def submit():
    """Submits approved application materials with user confirmation."""
    approved_materials = load_json_file(EDITED_MATERIALS_FILE)
    if not approved_materials:
        click.echo("No approved materials to submit. Run `export-docs` first.")
        logging.info("No approved materials to submit.")
        return

    for item in approved_materials:
        job_details = item.get('job_details', {})
        job_title = job_details.get('title', 'Unknown')
        company = job_details.get('company', {}).get('display_name', 'Unknown')

        click.echo(f"\n--- Job: {job_title} at {company} ---")
        confirm = click.confirm("Confirm submission? (This will send a mock email with attachments)", default=False)

        if confirm:
            success, message = submission.submit_application(item)
            if success:
                click.echo(f"Submission successful for {job_title}: {message}")
                logging.info(f"Submitted successfully: {job_title}")
            else:
                click.echo(f"Submission failed for {job_title}: {message}")
                logging.error(f"Submission failed for {job_title}: {message}")
        else:
            click.echo(f"Skipped submission for {job_title}")
            logging.info(f"Skipped submission for {job_title}")

@cli.command()
@click.argument('resume_path', type=click.Path(exists=True))
def update_profile(resume_path):
    """Parses a resume and uses AI to update the profile.json."""
    logging.info(f"Starting profile update with resume: {resume_path}")
    resume_text = parse_resume(resume_path)
    if not resume_text:
        return

    enhanced_data = enhance_profile_with_gemini(resume_text)
    if not enhanced_data:
        return

    profile = load_profile()
    profile.update(enhanced_data)
    
    original_skills = set(profile.get('skills', []))
    enhanced_skills = set(enhanced_data.get('enhanced_skills', []))
    profile['skills'] = sorted(list(original_skills.union(enhanced_skills)))
    
    if save_profile(profile):
        logging.info("Profile has been successfully updated with AI enhancements.")
    else:
        logging.error("Failed to save the updated profile. Please ensure your encryption key is set.")

@cli.command()
@click.option('--skill', multiple=True, help='Add a skill to your profile.')
@click.option('--location', help='Set your preferred job location.')
@click.option('--industry', help='Set your preferred industry.')
@click.option('--work-type', help='Set your preferred work type (e.g., full_time, contract).')
@click.option('--salary-range', help='Set your desired salary range (e.g., "100000-120000").')
def manual_update(skill, location, industry, work_type, salary_range):
    """Manually update your profile."""
    profile = load_profile()
    updated = False
    if skill:
        current_skills = set(profile.get('skills', []))
        current_skills.update(skill)
        profile['skills'] = sorted(list(current_skills))
        updated = True
    if location:
        profile['location'] = location
        updated = True
    if industry:
        profile['industry'] = industry
        updated = True
    if work_type:
        profile['work_type'] = work_type
        updated = True
    if salary_range:
        profile['salary_range'] = salary_range
        updated = True
        
    if updated:
        if save_profile(profile):
            logging.info("Profile updated successfully.")
        else:
            logging.error("Failed to save the updated profile.")
    else:
        logging.warning("No updates provided.")

@cli.command()
def generate_key():
    """Generates a new encryption key and prints it."""
    key = Fernet.generate_key()
    click.echo("Generated Encryption Key (add this to your .env file):\n")
    click.echo(f"ENCRYPTION_KEY={key.decode()}")

if __name__ == "__main__":
    cli()