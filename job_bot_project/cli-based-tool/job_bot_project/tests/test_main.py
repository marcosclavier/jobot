import pytest
from unittest.mock import patch, mock_open
import json
import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Assuming main.py is in the parent directory for imports
import sys
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
from click.testing import CliRunner
from main import (
    load_profile, save_profile, generate_application_materials, export_docs,
    PROFILE_FILE, PROFILE_HASH_FILE, GENERATED_MATERIALS_FILE, EDITED_MATERIALS_FILE
)

# Mock the Fernet class for encryption/decryption
class MockFernet:
    def __init__(self, key):
        pass

    def encrypt(self, data):
        return f"encrypted_{data.decode()}".encode()

    def decrypt(self, data):
        return data.decode().replace("encrypted_", "").encode()

@pytest.fixture
def mock_env_vars():
    with patch('os.getenv', return_value='test_key_123456789012345678901234567890='):
        yield

@pytest.fixture
def mock_fernet():
    with patch('main.Fernet', MockFernet):
        yield

@pytest.fixture
def mock_file_operations():
    with patch('builtins.open', new_callable=mock_open) as mock_file:
        yield mock_file

@pytest.fixture
def mock_json_load():
    with patch('json.load') as mock_load:
        yield mock_load

@pytest.fixture
def mock_json_dump():
    with patch('json.dump') as mock_dump:
        yield mock_dump

@pytest.fixture
def mock_os_path_exists():
    with patch('os.path.exists', return_value=True) as mock_exists:
        yield mock_exists

@pytest.fixture
def mock_generative_model():
    with patch('google.generativeai.GenerativeModel') as mock_model:
        mock_model.return_value.generate_content.return_value.text = json.dumps({
            "cover_letter": "This is a generated cover letter.",
            "resume_suggestions": ["Suggestion 1", "Suggestion 2"],
            "question_answers": [{"question": "Q1", "answer": "A1"}]
        })
        yield mock_model

@pytest.fixture
def mock_docx_document():
    with patch('docx.Document') as mock_doc:
        mock_instance = mock_doc.return_value
        mock_instance.add_heading.return_value = None
        mock_instance.add_paragraph.return_value.add_run.return_value = None
        mock_instance.add_paragraph.return_value.alignment = None
        mock_instance.sections = [
            type('Section', (object,), {'top_margin': None, 'bottom_margin': None, 'left_margin': None, 'right_margin': None})()
        ]
        yield mock_doc

# --- Tests for Profile Management ---
def test_load_profile_success(mock_env_vars, mock_fernet, mock_file_operations, mock_json_load):
    mock_file_operations.return_value.read.return_value = b'encrypted_{"skills": ["Python"], "location": "Remote"}'
    mock_json_load.return_value = {"skills": ["Python"], "location": "Remote"}
    profile = load_profile()
    assert profile == {"skills": ["Python"], "location": "Remote"}
    mock_file_operations.assert_called_with(PROFILE_FILE, 'rb')

def test_load_profile_file_not_found(mock_env_vars, mock_fernet, mock_file_operations):
    mock_file_operations.side_effect = FileNotFoundError
    profile = load_profile()
    assert profile == {}

def test_save_profile_success(mock_env_vars, mock_fernet, mock_file_operations, mock_json_dump, mock_os_path_exists):
    with patch('main.update_profile_hash') as mock_update_hash:
        profile_data = {"skills": ["Java"], "location": "New York"}
        result = save_profile(profile_data)
        assert result is True
        mock_file_operations.assert_called_with(PROFILE_FILE, 'wb')
        with patch('main.update_profile_hash') as mock_update_hash:
            profile_data = {"skills": ["Java"], "location": "New York"}
            result = save_profile(profile_data)
            assert result is True
            mock_file_operations.assert_called_with(PROFILE_FILE, 'wb')
            mock_json_dump.assert_called_once_with(profile_data, mock_file_operations(), indent=4)
        mock_update_hash.assert_called_once()

# --- Tests for AI-Powered Content Generation ---
def test_generate_application_materials_no_placeholders(mock_env_vars, mock_generative_model, mock_json_load):
    job_data = {
        "job_details": {
            "title": "Software Engineer",
            "full_description": "Develop software."
        }
    }
    profile = {"skills": ["Python"], "location": "Remote"}
    
    # Mock the generate_content response to ensure no placeholders are returned
    mock_generative_model.return_value.generate_content.return_value.text = json.dumps({
        "cover_letter": "This is a generated cover letter for Software Engineer.",
        "resume_suggestions": ["Update Python skills."],
        "question_answers": []
    })

    result = generate_application_materials(job_data, profile)
    assert result is not None
    assert "N/A" not in result['generated_materials']['cover_letter']
    assert "Applicant Name" not in result['generated_materials']['cover_letter']
    assert "No custom instructions provided." not in mock_generative_model.return_value.generate_content.call_args[0][0]

# --- Tests for Export Docs ---
def test_export_docs_no_placeholders_in_header(mock_env_vars, mock_file_operations, mock_json_load, mock_docx_document):
    # Mock load_profile to return an empty profile
    mock_json_load.side_effect = [
        {}, # For load_profile
        [{"job_details": {"title": "Dev", "company": {"display_name": "ABC"}}, "generated_materials": {"cover_letter": "CL", "refined_resume": "RES", "question_answers": []}}] # For load_json_file(EDITED_MATERIALS_FILE)
    ]
    
    # Mock os.makedirs to prevent actual directory creation
    runner = CliRunner()
    with patch('os.makedirs'), patch('main.add_styled_header') as mock_add_styled_header:
        result = runner.invoke(export_docs)
        assert result.exit_code == 0

        assert mock_add_styled_header.called
        # The second argument to add_styled_header is the text
        header_text_arg = mock_add_styled_header.call_args[0][1]
        assert "Applicant Name" not in header_text_arg
        assert "N/A" not in header_text_arg

def test_export_docs_with_profile_data(mock_env_vars, mock_file_operations, mock_json_load, mock_docx_document):
    # Mock load_profile to return a profile with data
    mock_json_load.side_effect = [
        {"name": "Test User", "contact_info": {"email": "test@example.com"}, "location": "Test City"}, # For load_profile
        [{"job_details": {"title": "Dev", "company": {"display_name": "ABC"}}, "generated_materials": {"cover_letter": "CL", "refined_resume": "RES", "question_answers": []}}] # For load_json_file(EDITED_MATERIALS_FILE)
    ]
    
    runner = CliRunner()
    with patch('os.makedirs'), patch('main.add_styled_header') as mock_add_styled_header:
        result = runner.invoke(export_docs)
        assert result.exit_code == 0

        assert mock_add_styled_header.called
        header_text_arg = mock_add_styled_header.call_args[0][1]
        assert "Test User" in header_text_arg
        assert "test@example.com" in header_text_arg
        assert "Test City" in header_text_arg
        assert "N/A" not in header_text_arg # Ensure N/A is not present if data exists

def test_export_docs_empty_materials_content(mock_env_vars, mock_file_operations, mock_json_load, mock_docx_document):
    # Mock load_profile and load_json_file to return empty materials
    mock_json_load.side_effect = [
        {}, # For load_profile
        [{"job_details": {"title": "Dev", "company": {"display_name": "ABC"}}, "generated_materials": {}}] # For load_json_file(EDITED_MATERIALS_FILE)
    ]

    runner = CliRunner()
    with patch('os.makedirs'), patch('main.add_formatted_content') as mock_add_formatted_content:
        result = runner.invoke(export_docs)
        assert result.exit_code == 0

        # Check that add_formatted_content is called with empty strings, not "Not generated."
        # It's called twice: once for cover letter, once for resume
        assert mock_add_formatted_content.call_count == 2
        assert mock_add_formatted_content.call_args_list[0][0][1] == ""
        assert mock_add_formatted_content.call_args_list[1][0][1] == ""
